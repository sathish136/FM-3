#!/usr/bin/env python3
"""Convert a filled .docx file to PDF using fpdf2 + python-docx.
Usage: python3 docx_to_pdf.py <input.docx> <output.pdf>
"""
import sys
from docx import Document
from fpdf import FPDF

def clean(text):
    if not text:
        return ""
    # Replace common unicode chars that Latin-1 can't handle
    replacements = {
        '\u2013': '-', '\u2014': '--', '\u2018': "'", '\u2019': "'",
        '\u201c': '"', '\u201d': '"', '\u2022': '*', '\u00a0': ' ',
        '\u00b3': '3', '\u00b2': '2', '\u00b9': '1',
    }
    for k, v in replacements.items():
        text = text.replace(k, v)
    return text.encode('latin-1', errors='replace').decode('latin-1')

def docx_to_pdf(input_path, output_path):
    doc = Document(input_path)
    pdf = FPDF(orientation='P', unit='mm', format='A4')
    pdf.set_auto_page_break(auto=True, margin=20)
    pdf.add_page()
    pdf.set_margins(20, 20, 20)

    def add_paragraph(text, bold=False, size=10, center=False):
        text = clean(text)
        if not text.strip():
            pdf.ln(3)
            return
        pdf.set_font('Helvetica', 'B' if bold else '', size)
        if center:
            pdf.multi_cell(0, 6, text, align='C')
        else:
            pdf.multi_cell(0, 6, text, align='L')
        pdf.ln(1)

    def add_table(table):
        col_count = max(len(row.cells) for row in table.rows) if table.rows else 0
        if col_count == 0:
            return
        page_w = pdf.w - pdf.l_margin - pdf.r_margin
        col_w = page_w / col_count
        for row in table.rows:
            texts = []
            for cell in row.cells:
                cell_text = ' '.join(p.text for p in cell.paragraphs).strip()
                texts.append(clean(cell_text))
            # check if it's a header row (all bold or first row)
            is_header = all(
                any(run.bold for run in cell.paragraphs[0].runs if cell.paragraphs)
                for cell in row.cells
            ) if row.cells else False
            pdf.set_font('Helvetica', 'B' if is_header else '', 8)
            row_h = 6
            for t in texts:
                pdf.cell(col_w, row_h, t[:60], border=1)
            pdf.ln()

    for element in doc.element.body:
        tag = element.tag.split('}')[-1] if '}' in element.tag else element.tag
        if tag == 'p':
            from docx.text.paragraph import Paragraph
            para = Paragraph(element, doc)
            style = para.style.name if para.style else ''
            is_heading = 'Heading' in style
            is_title = 'Title' in style
            size = 14 if is_title else (12 if is_heading else 10)
            bold = is_heading or is_title or all(r.bold for r in para.runs if r.text.strip())
            add_paragraph(para.text, bold=bold, size=size)
        elif tag == 'tbl':
            from docx.table import Table
            tbl = Table(element, doc)
            add_table(tbl)
            pdf.ln(3)

    pdf.output(output_path)

if __name__ == '__main__':
    if len(sys.argv) != 3:
        print("Usage: docx_to_pdf.py <input.docx> <output.pdf>", file=sys.stderr)
        sys.exit(1)
    docx_to_pdf(sys.argv[1], sys.argv[2])
    print("ok")
